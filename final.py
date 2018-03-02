#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from tornado.gen import coroutine
from tornado.ioloop import IOLoop
from tornado.httpclient import HTTPRequest, AsyncHTTPClient, HTTPError
from bs4 import BeautifulSoup
from docx import Document
from openpyxl import load_workbook

import time
import os
import re
import operator
import subprocess
import requests
import xlrd
import pymysql.cursors
# Define headers and starturl global

HEADERS = {'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8',
           'Accept-Encoding': 'gzip, deflate',
           'Accept-Language': 'en-US,en;q=0.9',
           'Cache-Control': 'max-age=0',
           'Connection': 'keep-alive',
           'Cookie': 'BIGipServerpool_ext_sgccecp=dLi+emekHSXw0p0qlHEOFwPG8f3ZUWy0UPTi92z2fDPzzz5qu84WJiSs0K5rXYZnsZCYog5OFFLXpCs=; WL_JSESSIONID=QVVnhk2KPTPgd7m4Bp918nBw1DWHmCr1nG4HL31Z2pB5P0c8f61Z!1877739744',
           'Host': 'ecp.sgcc.com.cn',
           'Referer': 'http://ecp.sgcc.com.cn/',
           'Upgrade-Insecure-Requests': '1',
           'User-Agent': 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/62.0.3202.94 Safari/537.36'
           }

# You may change the starturl below

starturl = 'http://ecp.sgcc.com.cn/project_list.jsp?site=global&column_code=014001001&project_type='

# 招标文件路径
tender_files_path = './zippp'
outpath = './docx'

connection = pymysql.connect(host='localhost',
                             user='root',
                             password='siliver88',
                             db='tender',
                             charset='utf8',
                             cursorclass=pymysql.cursors.DictCursor)

# Filter


class SpiderFilter(object):

    def __init__(self, starturl):
        self.starturl = starturl

    def PageUrl(self):
        # set the starurl range below
        for i in range(1, 11):
            yield self.starturl + str(i)

    # fetch the public bidding  urls
    def BiddingUrl(self):
        fileurls = []
        filenames = []
        fileresults = []
        # step 1 for yield black40s
        for i in self.PageUrl():
            tagurl_soup = BeautifulSoup(requests.get(i).text, 'lxml')
            tagurls = tagurl_soup.find_all(class_="black40")
            for url in tagurls:
                fileurls.append(url)
        # count the tds
        num = len(fileurls)

        # step 2 for filter '正在招标'
        for i in range(0, num, 4):
            if fileurls[i].string == '\r\n\t\t   \r\n\t\t   正在招标\r\n\t\t  \r\n\t\t    ':
                filenames.append([fileurls[i + 2].a.attrs['onclick'].strip("showProjectDetail")
                                  .strip("(").strip(");").split(",")[0].strip("'"),
                                  fileurls[i +
                                           2].a.attrs['onclick'].strip("showProjectDetail")
                                  .strip("(").strip(");").split(",")[1].strip("'")
                                  ])
        for filename in filenames:
            fileresults.append(
                'http://ecp.sgcc.com.cn/html/project/' + filename[0] + '/' + filename[1] + '.html')
        fileresults = set(fileresults)

        # 正在招标的链接保存在本地
        with open('./BiddingUrls.log', 'w') as fs:
            for i in fileresults:
                fs.write(i + '\n')
        print("Filtering bidding urls...")
        return fileresults

    # fetch  fileurls for downloadurls
    # list is faster than yield for asynchronously Fetching
    def FileUrl(self):
        downloadurls = []
        for i in self.BiddingUrl():
            preurl_soup = BeautifulSoup(requests.get(i).text, 'lxml')
            try:
                downloadname = preurl_soup.find(
                    href=re.compile("/structData/")).attrs['href']
            except ArrtributeError:
                continue
            finally:
                downloadurls.append('http://ecp.sgcc.com.cn' + downloadname)
        downloadurls = set(downloadurls)
        print("Hit " + str(len(downloadurls)) + " files")
        return downloadurls


# need BiddingUrls
class Det(object):

    def __init__(self):
        self.urls = []
        with open('./BiddingUrls.log', 'r') as f:
            file = f.readlines()
            for i in file:
                self.urls.append(i.strip('\n'))

    def Details(self, url, soup):
        details = {}
        details['公告页面'] = url
        for i in range(15):
            details['{}'.format(str(soup.find_all('tr')[i].find_all('td')[0].string).strip('：').strip(''))] = \
                str(soup.find_all('tr')[i].find_all(
                    'td')[1].string).strip('\r\n\t\t\t')

            # 如果下载链接为空， 则会报ArrtributeError
            filename = soup.find(href=re.compile("/structData/")).attrs['href']
            details['项目公告文件'] = 'http://ecp.sgcc.com.cn' + filename
            details['db_table_name'] = os.path.basename(
                filename).split('.')[0]
            ctime = soup.find(class_="articleTitle_details").string
            details['发布时间'] = ctime
        return details

    def GetSoup(self, url):
        resp = requests.get(url)
        # 解决中文乱码
        forbs4 = resp.text.encode(resp.encoding).decode('utf-8')
        soup = BeautifulSoup(forbs4, 'lxml')
        return soup

    def ff(self):
        # 内嵌dict 格式
        foo = []
        for url in self.urls:
            foo.append(self.Details(url, self.GetSoup(url)))
        result = {}
        for i in range(len(foo)):
            result[i] = foo[i]
        return result

# Spider get and download method
# Coroutine for downloading to safe time
# Reaplace the PageUrl() with FileUrl()
# to download files


class SpiderMenthod(object):

    def __init__(self):
        self.http = AsyncHTTPClient()
        if not os.path.exists(tender_files_path):
            os.mkdir(tender_files_path)
        if not os.path.exists(outpath):
            os.mkdir(outpath)
        else:
            print("Dirs are ready.")

    @coroutine
    def get(self, url):
        request = HTTPRequest(url=url,
                              method='GET',
                              headers=HEADERS,
                              connect_timeout=40.0,
                              request_timeout=40.0,
                              follow_redirects=False,
                              max_redirects=False,
                              )
        yield self.http.fetch(request, callback=self.find, raise_error=False)

    def find(self, response):
        if response.error:
            print(response.error)
        print(response.code, response.effective_url, response.request_time)
        with open(os.path.join(tender_files_path, os.path.basename(response.effective_url)), 'wb') as f:
            f.write(response.body)

    @coroutine
    def downloadUrls(self):
        urls = SpiderFilter(starturl).FileUrl()
        print('Fetching urls asynchronously...')
        t1 = time.time()
        yield [self.get(url) for url in urls]
        t = time.time() - t1
        print(t)

# 以下是解解压，写数据库部分
# Need  path first


class FilePath(object):

    def __init__(self, sourcepath, outpath):
        self.sourcepath = sourcepath
        self.outpath = outpath

    # Return full path list
    def WalkFile(self, path):
        for fpath, dirs, fs in os.walk(path):
            for f in fs:
                yield os.path.join(fpath, f)

    def UnzipFile(self):
        for filename in os.listdir(self.sourcepath):
            subprocess.call(['unar', '-o', str(self.outpath) + '/' +
                             filename.split('.')[0], str(self.sourcepath) + '/' + filename])

        for fpath, dirs, fn in os.walk(self.outpath):
            for f in fn:
                if f.endswith('.rar'):
                    subprocess.call(
                        ['unar', '-o', fpath, os.path.join(fpath, f)])

    def ConvertFiles(self):
        for fpath, dirs, fs in os.walk(self.outpath):
            for f in fs:
                if f.endswith('.doc'):
                    subprocess.call(['soffice', '--headless', '--convert-to',
                                     'docx', os.path.join(fpath, f), '--outdir', fpath])
                    subprocess.call(['rm', os.path.join(fpath, f)])
                #if f.endswith('.xlsx'):
                #    subprocess.call(['soffice', '--headless', '--convert-to',
                #                     'xls', os.path.join(fpath, f), '--outdir', fpath])
                #    subprocess.call(['rm', os.path.join(fpath, f)])
        print("Converted Done!")
        return

# ======================
# 以下是数据库操作

# 用于清空数据库所有表单


class DBClean(object):
    def __init__(self):

        try:
            with connection.cursor() as cursor:
                sql = '''SHOW TABLES'''
                cursor.execute(sql)
                dbtable_names_for_now = cursor.fetchall()
                for i in dbtable_names_for_now:
                    # print(i)
                    checksql = '''DROP TABLE IF EXISTS {}'''.format(
                        i['Tables_in_tender'])
                    cursor.execute(checksql)
                connection.commit()
        finally:
            print('All tables are deleted!')


# 写入数据
class Write_To_DB(object):
    def __init__(self, path):
        
        # 用于接收outpath
        self.path = path
        self.__Create_tables()
        self.__Write_data()
        self.__Write_files()

    # The method for dealing with single file
    # If you wanna deal with a list , use `for`

    def __Create_tables(self):
        tables = []
        bigresult = Det().ff()
        for i in range(len(bigresult)):
            tables.append('zb_' + bigresult[i]['db_table_name'])

        # 检查抓取的文件名是否在数据库中
        # 如果存在，则不写入;否则写入数据库
        dbtable_names = []
        dbtables_not_written = []

        # 获取数据库中所有表名
        try:
            with connection.cursor() as cursor:
                sql = '''SHOW TABLES'''
                cursor.execute(sql)
                # for_now 为临时变量
                dbtable_names_for_now = cursor.fetchall()
                for i in range(len(dbtable_names_for_now)):
                    dbtable_names.append(
                        dbtable_names_for_now[i]['Tables_in_tender'])

            if operator.eq(tables, dbtable_names):
                pass
            else:
                for i in tables:
                    if i not in dbtable_names:
                        dbtables_not_written.append(i)

            # 测试dbtables_not_written 是否奏效
            # 如果数据库已经写入，则返回空列表
            # print(len(dbtables_not_written))

            if len(dbtables_not_written) == 0:
                pass
            else:
                # 建立table
                with connection.cursor() as cursor:
                    for name in dbtables_not_written:
                        createtablesql = '''CREATE TABLE {}(
                                  `id` int(10) NOT NULL AUTO_INCREMENT COMMENT '记录id',
                                  `isAlive` VarChar(15) NOT NULL DEFAULT '' COMMENT '项目状态',
                                  `zb_url` VarChar(255) DEFAULT '' COMMENT '公告页面',
                                  `file_url` VarChar(255) NOT NULL DEFAULT '' COMMENT '项目公告文件',
                                  `title` VarChar(255) NOT NULL DEFAULT '' COMMENT '项目名称',
                                  `content` MediumText COMMENT '段落内容',
                                  `lettercard` MediumText COMMENT '单元格内容',
                                  `ctime` VarChar(30) NOT NULL DEFAULT 0 COMMENT '发布时间',
                                  `deadline` VarChar(30) NOT NULL DEFAULT 0 COMMENT '截标时间',
                                  `bsdeadline` VarChar(30) NOT NULL DEFAULT 0 COMMENT '标书购买截止时间',
                                  `projectcode` VarChar(50) DEFAULT '' COMMENT '项目编号',
                                  `kbsj` VarChar(30) NOT NULL DEFAULT 0 COMMENT '开标时间',
                                  `bsfyxs` VarChar(100) NOT NULL DEFAULT '下载前支付' COMMENT '标书费用形式',
                                  `xmjs` VarChar(100) DEFAULT '' COMMENT '项目介绍',
                                  `telephone` VarChar(50) DEFAULT '' COMMENT '联系电话',
                                  `projectclass` VarChar(255) NOT NULL DEFAULT '' COMMENT '项目类型',
                                  `tenderee` VarChar(100) NOT NULL DEFAULT '' COMMENT '招标人',
                                  `contact` VarChar(100) DEFAULT '' COMMENT '联系人',
                                  `agent` VarChar(255) NOT NULL DEFAULT '' COMMENT '代理机构',
                                  `dbtb` VarChar(50) NOT NULL DEFAULT '' COMMENT 'db_table_name',
                                  PRIMARY KEY(`id`)
                                )ENGINE=InnoDB DEFAULT CHARSET=utf8'''.format(name)
                        cursor.execute(createtablesql)
                connection.commit()
        finally:
            print("Tables are created!")

    def __Write_data(self):
        bigdata = Det().ff()

        try:
            with connection.cursor() as cursor:
                for i in range(len(bigdata)):
                    tabname = 'zb_' + bigdata[i]['db_table_name']
                    isAlive = bigdata[i]['项目状态']
                    zb_url = bigdata[i]['公告页面']
                    file_url = bigdata[i]['项目公告文件']
                    title = bigdata[i]['项目名称']
                    ctime = bigdata[i]['发布时间']
                    deadline = bigdata[i]['截标时间']
                    bsdeadline = bigdata[i]['标书购买截止时间']
                    projectcode = bigdata[i]['项目编号']
                    kbsj = bigdata[i]['开标时间']
                    bsfyxs = bigdata[i]['标书费用形式']
                    xmjs = bigdata[i]['项目介绍']
                    telephone = bigdata[i]['联系电话']
                    projectclass = bigdata[i]['项目类型']
                    tenderee = bigdata[i]['招标人']
                    contact = bigdata[i]['联系人']
                    agent = bigdata[i]['代理机构']
                    dbtb = 'zb_' + bigdata[i]['db_table_name']
                    content = '段落: '
                    lettercard = '单元格: '

                    insertsql = '''INSERT INTO {} (isAlive, zb_url, file_url, title, ctime,
                                deadline, bsdeadline, projectcode, kbsj, bsfyxs, xmjs, telephone,
                                projectclass, tenderee, contact, agent, dbtb, content, lettercard) VALUES \
                               ('{}','{}','{}','{}','{}','{}','{}','{}','{}','{}','{}','{}','{}','{}','{}','{}','{}','{}','{}')''' \
                                .format(tabname, isAlive, zb_url, file_url, title, ctime,
                                        deadline, bsdeadline, projectcode, kbsj, bsfyxs, xmjs, telephone,
                                        projectclass, tenderee, contact, agent, dbtb, content, lettercard)
                    cursor.execute(insertsql)
                connection.commit()
        finally:
            # 注意，此处connection 只为模块测试使用
            # 正式环境在所有数据库操作结束后再写close()
            print("Data is inserted!")

    # 将docx, xls 文件写入数据库
    def WalkFile(self, path):
        for fpath, dirs, fs in os.walk(path):
            for f in fs:
                yield os.path.join(fpath,f)
            
    def Write_file_To_DB(self, tbname, path):

        for name in self.WalkFile(outpath):
            if os.path.basename(name).startswith('~$'):
                print("Invalid file name :", name)
                subprocess.call(['rm', name])

        content = []
        lettercard = []
            
        if path.endswith('.docx'):
            print("docx files... ", path)
            document = Document(path)
            for para in document.paragraphs:
                content.append(para.text)

            tempTable = document.tables
            for i in range(len(tempTable)):
                table = tempTable[i]
                for x in table.rows:
                    try:
                        for y in x.cells:
                            # use extend() for safe
                            lettercard.extend(y.text)
                    except IndexError:
                        continue
                    finally:
                        print('docx 文档中有table 存在未对齐情况')

        if path.endswith('.xls'):
            print("xls files...", path)
            book = xlrd.open_workbook(path)
            worksheets = book.sheets()
            for num in range(len(worksheets)):
                worksheetn = book.sheet_by_index(num)
                num_rows = worksheetn.nrows
                for curr_row in range(num_rows):
                    row = worksheetn.row_values(curr_row)
                    if row != None: 
                        lettercard.extend(str(row))

        if path.endswith('.xlsx'):
            print("xlsx files...", path)
            wb = load_workbook(path) 
            sheetnamelist = wb.get_sheet_names()
            for name in sheetnamelist:
                sheet = wb.get_sheet_by_name(name)
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value != None:
                            lettercard.extend(str(cell.value))

        while '' in content:
            content.remove('')
        
        while '' in lettercard:
            lettercard.remove('')        

        content = ','.join(content)
        lettercard = ''.join(lettercard)

        try:
            with connection.cursor() as cursor:
                sql = '''UPDATE {} SET content=CONCAT(content, %s),lettercard=CONCAT(lettercard,%s) WHERE id=1'''.format(tbname)
                cursor.execute(sql,(content,lettercard))
            connection.commit()
        finally:
            print("Updating table...")
    
    def __Write_files(self):
        for i in os.listdir(self.path):
            for x in self.WalkFile(outpath + '/' + i):
                self.Write_file_To_DB('zb_' + i, x)
        print('All files are written into DB.')

def main():
    # 先抓数据
    DBClean()
    spider = SpiderMenthod()
    loop = IOLoop.current()
    loop.run_sync(spider.downloadUrls)
    fp = FilePath(tender_files_path, outpath)
    fp.UnzipFile()
    fp.ConvertFiles()
    Write_To_DB(outpath)
    connection.close()

if __name__ == "__main__":
    main()
